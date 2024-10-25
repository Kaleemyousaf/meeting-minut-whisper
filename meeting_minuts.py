import streamlit as st
import whisper
from docx import Document
from datetime import datetime

# Initialize a list to hold pre-added members and their transcriptions
members = []
transcriptions = []  # Changed from dict to list to maintain order

# Initialize the Whisper model
model = whisper.load_model("medium.en")

# Function to transcribe audio
def transcribe_audio(selected_members, audio):
    if not selected_members:
        return "No member selected."

    if not audio:
        return "No audio file provided."

    # Load and process the audio
    audio_data = whisper.load_audio(audio)
    result = model.transcribe(audio_data)

    # Store the transcription in order with the member's data
    transcription = result["text"]
    for member in selected_members:
        transcriptions.append((member, transcription))

    return f"Transcription for {', '.join(selected_members)}: {transcription}"

# Function to display meeting minutes
def display_minutes(meeting_title, date, venue, requested_by, agenda, action_items, next_meeting, additional_notes):
    minutes = f"{meeting_title}\n\n"
    minutes += f"Date and Time: {date}\n"
    minutes += f"Venue: {venue}\n"
    minutes += f"Requested by: {requested_by}\n\n"
    minutes += "Meeting Attendees:\n"
    for member in members:
        minutes += f"  - {member}\n"
    minutes += "\nAgenda:\n"
    minutes += f"{agenda}\n\n"
    minutes += "Discussion:\n"
    for member, transcription in transcriptions:
        minutes += f"  - {member}: {transcription}\n"
    minutes += "\nAction Items:\n"
    minutes += f"{action_items}\n\n"
    minutes += "Next Meeting:\n"
    minutes += f"{next_meeting}\n\n"
    minutes += "Additional Notes:\n"
    minutes += f"{additional_notes}\n\n"

    return minutes

# Function to save meeting minutes as a Word document
def save_minutes(minutes):
    try:
        doc = Document()
        doc.add_heading("Meeting Minutes", 0)

        for line in minutes.split("\n"):
            if line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)

        file_name = "Meeting_Minutes.docx"
        doc.save(file_name)
        return file_name

    except Exception as e:
        return f"An error occurred while saving the document: {e}"

# Streamlit interface
st.title("Meeting Minutes and Transcription App")

# Input fields for meeting details
meeting_title = st.text_input("Meeting Title", placeholder="Enter the meeting title...")
date = st.text_input("Select Meeting Date", value=datetime.now().strftime('%Y-%m-%d'))
venue = st.text_input("Venue", placeholder="Enter the venue...")
requested_by = st.text_input("Requested by", placeholder="Enter the name of the person who requested the meeting...")

# Section to add/select members
new_member = st.text_input("Add New Member")
if st.button("Add Member"):
    if new_member and new_member not in members:
        members.append(new_member)

selected_members = st.multiselect("Select Members", members)

# Audio upload for transcription
audio = st.file_uploader("Upload Member's Audio", type=["wav", "mp3", "m4a"])

# Text inputs for agenda, action items, next meeting, and additional notes
agenda = st.text_area("Agenda", placeholder="List the agenda points here...")
action_items = st.text_area("Action Items", placeholder="List the action items here...")
next_meeting = st.text_input("Next Meeting", placeholder="Details for the next meeting...")
additional_notes = st.text_area("Additional Notes", placeholder="Any additional notes...")

# Buttons to transcribe and display minutes
if st.button("Transcribe"):
    transcription_result = transcribe_audio(selected_members, audio)
    st.write(transcription_result)

if st.button("Display Meeting Minutes"):
    minutes = display_minutes(meeting_title, date, venue, requested_by, agenda, action_items, next_meeting, additional_notes)
    st.text_area("Meeting Minutes", value=minutes, height=400)

# Button to save the meeting minutes as a Word document
if st.button("Save Meeting Minutes"):
    minutes_text = display_minutes(meeting_title, date, venue, requested_by, agenda, action_items, next_meeting, additional_notes)
    file_path = save_minutes(minutes_text)
    with open(file_path, "rb") as file:
        st.download_button("Download Meeting Minutes Document", data=file, file_name="Meeting_Minutes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
